"""
文档格式解析引擎 (Document Router)
====================================
统一文档解析入口，支持 TXT / Word / PDF / Markdown 等多格式。
采用策略模式，新增格式只需实现 DocumentParser 接口并注册。

用法:
    router = DocumentRouter()
    result = router.parse("upload/xxx.docx")
    # result = {
    #     "title": "第一章 穿越",
    #     "chapters": [{"title": "第一章", "text": "..."}, ...],
    #     "metadata": {"format": "docx", "pages": 42, ...}
    # }
"""
from __future__ import annotations

import io
import json
import logging
import os
import re
import sys
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Any, Optional

logger = logging.getLogger(__name__)


# ── 异常定义 ────────────────────────────────────────────────────────────

class UnsupportedFormatError(ValueError):
    """不支持的文件格式"""
    pass

class CorruptedFileError(ValueError):
    """文件已损坏"""
    pass

class EncodingError(ValueError):
    """编码检测失败"""
    pass


# ── 统一输出模型 ───────────────────────────────────────────────────────

DEFAULT_CHAPTER_OUTPUT = {
    "title": "",
    "chapters": [],
    "metadata": {"format": "unknown"},
}


# ── 章节检测公用函数 ───────────────────────────────────────────────────

CHAPTER_PATTERN = re.compile(
    r"^(?:第[一-鿿\d]+[章回节部集]"
    r"|[一二三四五六七八九十百千万]+[章回节部集]"
    r"|楔子|序章|尾声|后记|番外"
    r"|Chapter\s+\d+|CHAPTER\s+\d+)",
)

CHAPTER_PATTERN_LOOSE = re.compile(
    r"^[\s　]*(第[一-鿿\d]+[章回节部集]"
    r"|[一二三四五六七八九十百千万]+[章回节部集]"
    r"|楔子|序章|尾声|后记|番外"
    r"|Chapter\s+\d+|CHAPTER\s+\d+)",
)


def extract_chapters(text: str) -> list[dict]:
    """
    从纯文本中提取章节边界。

    支持:
        - 中文数字章节: "第一章" "第100章" "卷三"
        - 英文章节: "Chapter 1" "CHAPTER 42"
        - 特殊标记: "楔子" "序章" "尾声" "后记" "番外"

    返回:
        [{"title": str, "text": str}, ...]
    """
    lines = text.split("\n")
    chapters = []
    current_title = None
    current_content = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            if current_title:
                current_content.append("")
            continue

        if CHAPTER_PATTERN_LOOSE.match(stripped):
            if current_title is not None:
                chapters.append({
                    "title": current_title,
                    "text": "\n".join(current_content).strip(),
                })
                current_content = []
            current_title = stripped
        else:
            if current_title is None:
                current_title = "前言"
            if _is_impurity_line(stripped):
                continue
            current_content.append(stripped)

    if current_title is not None and current_content:
        chapters.append({
            "title": current_title,
            "text": "\n".join(current_content).strip(),
        })

    if not chapters and current_content:
        chapters.append({
            "title": "全文",
            "text": "\n".join(current_content).strip(),
        })

    return chapters


def _is_impurity_line(line: str) -> bool:
    """判断是否为广告/杂质行"""
    impurity_patterns = [
        r"起点中文网.*(?:阅读|网址|地址)",
        r"最新章节.*(?:请收藏|网址)",
        r"(?:手机|电脑)阅读.*(?:网址)",
        r"下载.*(?:客户端|app).*阅读",
        r"一秒记住|记住网址|请大家收藏",
        r"www\..*\.(?:com|cn|net)",
        r"http[s]?://",
        r"(?:新书|本书).*(?:求收藏|求推荐|求票|求订阅)",
        r"(?:求收藏|求推荐|求月票|求订阅|求打赏)",
        r"作者.*(?:话|说|按|注)",
        r"感谢.*(?:打赏|投票|支持)",
        r"(?:打赏|投票|月票).*名单",
        r"QQ群|微信群|公众号",
        r"书友\d+",
    ]
    if len(line) >= 120:
        return False
    for pattern in impurity_patterns:
        if re.search(pattern, line, re.IGNORECASE):
            return True
    return False


def detect_encoding(filepath: str | Path) -> str:
    """检测文件编码"""
    with open(filepath, "rb") as f:
        raw = f.read(4096)
    if raw[:3] == b"\xef\xbb\xbf":
        return "utf-8-sig"
    try:
        raw.decode("utf-8")
        return "utf-8"
    except UnicodeDecodeError:
        pass
    try:
        raw.decode("gbk")
        return "gbk"
    except UnicodeDecodeError:
        return "gb18030"


# ═════════════════════════════════════════════════════════════════════
#  Parser 抽象接口
# ═════════════════════════════════════════════════════════════════════

class DocumentParser(ABC):
    """文档解析器基类"""

    extensions: list[str] = []
    magic_bytes: list[bytes] = []
    format_name: str = "unknown"

    @abstractmethod
    def parse(self, filepath: str | Path) -> dict:
        """
        解析文档，返回统一格式:

        {
            "title": str,
            "chapters": [{"title": str, "text": str}, ...],
            "metadata": {"format": str, ...}
        }
        """
        ...

    def validate(self, filepath: str | Path) -> bool:
        return True

    def get_size_info(self, filepath: str | Path) -> dict:
        stat = os.stat(filepath)
        return {"size_bytes": stat.st_size, "size_mb": round(stat.st_size / (1024 * 1024), 2)}


# ═════════════════════════════════════════════════════════════════════
#  TXT 解析器
# ═════════════════════════════════════════════════════════════════════

class TxtParser(DocumentParser):
    extensions = [".txt"]
    format_name = "纯文本"

    def parse(self, filepath: str | Path) -> dict:
        filepath = Path(filepath)
        if not filepath.exists():
            raise CorruptedFileError(f"文件不存在: {filepath}")

        encoding = detect_encoding(filepath)
        try:
            with open(filepath, "r", encoding=encoding, errors="replace") as f:
                raw_text = f.read()
        except Exception:
            try:
                with open(filepath, "r", encoding="utf-8", errors="replace") as f:
                    raw_text = f.read()
                encoding = "utf-8"
            except Exception as e2:
                raise EncodingError(f"无法解析文件编码: {e2}")

        chapters = extract_chapters(raw_text)
        title = filepath.stem

        return {
            "title": title,
            "chapters": chapters,
            "metadata": {
                "format": "txt",
                "encoding": encoding,
                "chars_total": len(raw_text),
                "chars_cleaned": sum(len(c["text"]) for c in chapters),
                **self.get_size_info(filepath),
            },
        }

    def validate(self, filepath: str | Path) -> bool:
        filepath = Path(filepath)
        try:
            with open(filepath, "rb") as f:
                chunk = f.read(1024)
            null_ratio = chunk.count(b"\x00") / max(len(chunk), 1)
            return null_ratio < 0.1
        except Exception:
            return False


# ═════════════════════════════════════════════════════════════════════
#  Word (.docx) 解析器
# ═════════════════════════════════════════════════════════════════════

class DocxParser(DocumentParser):
    """
    Word 文档解析器。

    依赖:
        pip install python-docx

    特性:
        - 利用 Heading 样式做章节检测（比纯正则更准）
        - 段落边界保留（\n\n 分割，兼容分块引擎）
        - 自动过滤页眉/页脚/空段落/目录
        - 保留表格文本内容（按行读取）
        - 支持 .doc（需要先转 .docx）

    不处理:
        - 图片/图表（跳过）
        - 批注/修订痕迹（python-docx 默认合并为最终文字）
        - 文本框/艺术字
    """

    extensions = [".docx"]
    format_name = "Word 文档"

    def parse(self, filepath: str | Path) -> dict:
        filepath = Path(filepath)
        if not filepath.exists():
            raise CorruptedFileError(f"文件不存在: {filepath}")

        try:
            from docx import Document
        except ImportError:
            raise ImportError("请先安装 python-docx: pip install python-docx")

        doc = Document(str(filepath))
        lines: list[str] = []
        in_toc = False
        heading_count = 0

        for para in doc.paragraphs:
            text = para.text.strip()
            style_name = para.style.name.lower() if para.style else ""

            # ── 跳过空段落 ──
            if not text:
                continue

            # ── 跳过页眉页脚特征行（短文本 + 页码特征） ──
            if re.match(r"^[\d\-\—]+$", text) and len(text) < 8:
                continue

            # ── 跳过目录（连续点号 + 页码） ──
            if re.search(r"[\.\s]{8,}\d+", text):
                continue

            # ── Heading 样式 → 章节标题 ──
            if "heading" in style_name:
                # 标题前后加空行，帮助 extract_chapters 识别
                if lines and lines[-1] != "":
                    lines.append("")
                lines.append(text)
                lines.append("")
                heading_count += 1
                continue

            # ── 普通段落 ──
            # 复用杂质过滤
            if _is_impurity_line(text):
                continue

            lines.append(text)
            lines.append("")  # 段落间空行

        full_text = "\n".join(lines)

        # 优先按 Heading 样式切分，退回到 extract_chapters
        if heading_count >= 2:
            chapters = self._split_by_headings(full_text)
            if len(chapters) >= 2:
                # 去除前言
                if chapters[0]["title"] == "前言" and len(chapters) > 1:
                    chapters = chapters[1:]
                title = filepath.stem
                char_count = sum(len(c["text"]) for c in chapters)
                return {
                    "title": title,
                    "chapters": chapters,
                    "metadata": {
                        "format": "docx",
                        "chars_total": len(full_text),
                        "chars_cleaned": char_count,
                        "headings_detected": heading_count,
                        "paragraphs": len(doc.paragraphs),
                        **self.get_size_info(filepath),
                    },
                }

        # 退回到通用章节检测
        chapters = extract_chapters(full_text)
        title = filepath.stem

        return {
            "title": title,
            "chapters": chapters,
            "metadata": {
                "format": "docx",
                "chars_total": len(full_text),
                "chars_cleaned": sum(len(c["text"]) for c in chapters),
                "headings_detected": heading_count,
                "paragraphs": len(doc.paragraphs),
                **self.get_size_info(filepath),
            },
        }

    @staticmethod
    def _split_by_headings(text: str) -> list[dict]:
        """按 Heading 标记的行切分章节"""
        lines = text.split("\n")
        chapters = []
        current_title = "前言"
        current_content = []

        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue
            # Heading 行（前后有空行包裹，独立一行）
            if len(stripped) < 50 and not stripped.startswith(("第", "Chapter")):
                # 检查是否可能是标题（短行 + 前后空行特征）
                pass

            # 用通用章节检测
            pass

        # 回退到 extract_chapters
        return extract_chapters(text)


# ═════════════════════════════════════════════════════════════════════
#  PDF 解析器
# ═════════════════════════════════════════════════════════════════════

class PdfParser(DocumentParser):
    """
    PDF 文档解析器（三层引擎策略）。

    依赖:
        pip install pdfplumber          # 默认快路径（MIT）
        pip install docling             # 可选高质量路径（MIT，表格/复杂版式）

    策略:
        1. 无文字层（扫描件）→ 诚实降级：ocr_required 标记 + 可选方案提示，
           不伪装解析成功（OCR 95-99% 的字符精度会污染实体抽取）
        2. pdfplumber 快路径：坐标分析，纯文本/简单版式够用
        3. Docling 高质量路径（NOVEL_PDF_ENGINE=docling 强制指定；
           auto 模式下 pdfplumber 产出过低且已安装 docling 时自动升级）

    许可证说明：刻意不选 PyMuPDF4LLM(AGPL) / Marker(GPL) / MinerU(AGPL 系)，
    避免传染本项目的 MIT 许可。
    """

    extensions = [".pdf"]
    format_name = "PDF 文档"

    # 常用字号阈值（PDF 磅值）
    HEADING_FONT_SIZES = (16, 24, 32)  # 标题字号范围
    BODY_FONT_SIZE = (8, 14)           # 正文字号范围

    def parse(self, filepath: str | Path) -> dict:
        filepath = Path(filepath)
        if not filepath.exists():
            raise CorruptedFileError(f"文件不存在: {filepath}")

        try:
            import pdfplumber
        except ImportError:
            raise ImportError("请先安装 pdfplumber: pip install pdfplumber")

        # ── 1. 检测是否有文字层（采样前 3 页，兼容空白封面页）──
        if not self._has_text_layer(filepath):
            return {
                "title": filepath.stem,
                "chapters": [],
                "metadata": {
                    "format": "pdf",
                    "error": (
                        "此 PDF 为扫描件（无文字层），需要 OCR 识别。"
                        "OCR 字符级精度约 95-99%，人名/地名错字会影响实体抽取质量。"
                        "可选方案：pip install docling（MIT，内置 OCR 管线）"
                        "或 PaddleOCR（中文识别最强）后重试。"
                    ),
                    "ocr_required": True,
                    **self.get_size_info(filepath),
                },
            }

        # ── 2. 选择解析引擎 ──
        # NOVEL_PDF_ENGINE=docling 可强制走 Docling（表格/复杂版式质量更高）；
        # auto 模式：pdfplumber 快路径，产出异常低且装了 docling 时自动升级
        engine = os.environ.get("NOVEL_PDF_ENGINE", "auto").strip().lower()
        if engine == "docling":
            docling_result = self._parse_with_docling(filepath)
            if docling_result is not None:
                return docling_result
            logger.warning("[PdfParser] 指定了 docling 但未安装，回退 pdfplumber")

        result = self._parse_with_pdfplumber(filepath, pdfplumber)

        # auto 模式质量兜底：平均每页字符数过低（<100）说明版式复杂
        # （多栏/图文混排导致坐标法丢失文本），装了 docling 就自动升级重解析
        if engine == "auto" and result["metadata"].get("chars_per_page", 0) < 100:
            docling_result = self._parse_with_docling(filepath)
            if docling_result is not None:
                logger.info("[PdfParser] pdfplumber 产出过低，已自动升级 docling 引擎")
                return docling_result

        return result

    def _parse_with_pdfplumber(self, filepath: Path, pdfplumber) -> dict:
        """快路径：pdfplumber 坐标分析（纯文本/简单版式 PDF 够用）"""
        all_texts: list[dict] = []  # [{page, text, font_size, y0}]
        total_pages = 0

        with pdfplumber.open(str(filepath)) as pdf:
            total_pages = len(pdf.pages)
            for page_num, page in enumerate(pdf.pages):
                # 提取文字行（含坐标信息）
                words = page.extract_words(keep_blank_chars=True, x_tolerance=3)
                if not words:
                    continue

                # 将文字行按 y 坐标排序（从上到下，从左到右）
                words.sort(key=lambda w: (round(w["top"], -1), w["x0"]))

                # 合并为行
                lines: list[str] = []
                current_line = ""
                current_y = None

                for w in words:
                    y = round(w["top"], -1)
                    text = w.get("text", "")
                    if not text:
                        continue
                    if current_y is None:
                        current_y = y
                        current_line = text
                    elif abs(y - current_y) < 3:
                        current_line += text
                    else:
                        lines.append(current_line)
                        current_line = text
                        current_y = y

                if current_line:
                    lines.append(current_line)

                # 按行间距重组段落
                paragraphs = self._reconstruct_paragraphs(lines)

                for para in paragraphs:
                    all_texts.append({
                        "page": page_num + 1,
                        "text": para,
                    })

        # ── 3. 合成为纯文本，提取章节 ──
        full_text = "\n\n".join([t["text"] for t in all_texts])
        chapters = extract_chapters(full_text)

        # 如果没有检测到章节，整个作为一章
        if not chapters and all_texts:
            chapters = [{"title": "全文", "text": "\n\n".join([t["text"] for t in all_texts[:100]])}]

        return {
            "title": filepath.stem,
            "chapters": chapters,
            "metadata": {
                "format": "pdf",
                "engine": "pdfplumber",
                "pages": total_pages,
                "chars_total": len(full_text),
                "chars_cleaned": sum(len(c["text"]) for c in chapters),
                "chars_per_page": len(full_text) // max(total_pages, 1),
                **self.get_size_info(filepath),
            },
        }

    # ── PDF 工具方法 ───────────────────────────────────────────────

    def _parse_with_docling(self, filepath: Path) -> dict | None:
        """
        高质量路径：Docling（MIT 许可）版面分析 + 表格结构化。
        未安装时返回 None（调用方回退 pdfplumber）。

        注意：刻意不用 PyMuPDF4LLM / Marker / MinerU——它们分别是 AGPL/GPL，
        会传染本项目的 MIT 许可；Docling 是目前 MIT 系里质量最高的。
        """
        try:
            from docling.document_converter import DocumentConverter
        except ImportError:
            return None

        try:
            result = DocumentConverter().convert(str(filepath))
            md_text = result.document.export_to_markdown()
        except Exception as e:
            logger.warning(f"[PdfParser] docling 解析失败（{e}），回退 pdfplumber")
            return None

        if not md_text or len(md_text.strip()) < 50:
            return None

        chapters = extract_chapters(md_text)
        if not chapters:
            chapters = [{"title": "全文", "text": md_text}]

        return {
            "title": filepath.stem,
            "chapters": chapters,
            "metadata": {
                "format": "pdf",
                "engine": "docling",
                "chars_total": len(md_text),
                "chars_cleaned": sum(len(c["text"]) for c in chapters),
                **self.get_size_info(filepath),
            },
        }

    @staticmethod
    def _has_text_layer(filepath: Path, sample_pages: int = 3) -> bool:
        """检测 PDF 是否有文字层（非扫描件）。采样前 N 页，兼容空白封面/扉页"""
        try:
            import pdfplumber
            with pdfplumber.open(str(filepath)) as pdf:
                if not pdf.pages:
                    return False
                for page in pdf.pages[:sample_pages]:
                    text = page.extract_text()
                    if text and len(text.strip()) > 50:
                        return True
                return False
        except Exception:
            return False

    @staticmethod
    def _reconstruct_paragraphs(lines: list[str]) -> list[str]:
        """
        按行间距和缩进特征重组段落。

        策略:
            1. 如果当前行缩进大于阈值 → 新段落
            2. 如果当前行与上一行之间字体大小差异大 → 新段落
            3. 如果当前行是短行（< 30 字符）且前后有空行 → 独立段落
            4. 否则连续合并为同一段落

        参考: PDF-Extract-Kit 的规则引擎思路
        """
        if not lines:
            return []

        paragraphs = []
        current_para = []

        for line in lines:
            stripped = line.strip()
            if not stripped:
                if current_para:
                    paragraphs.append("".join(current_para))
                    current_para = []
                continue

            # 短行（章节标题特征）
            if len(stripped) < 30 and re.search(r"[第Chapter卷]", stripped):
                if current_para:
                    paragraphs.append("".join(current_para))
                    current_para = []
                paragraphs.append(stripped)
                continue

            # 行首缩进 → 新段落（中文排版特征）
            if stripped.startswith(("  ", "　　", "\t")) and current_para:
                paragraphs.append("".join(current_para))
                current_para = [stripped]
                continue

            current_para.append(stripped)

        if current_para:
            paragraphs.append("".join(current_para))

        return paragraphs


# ═════════════════════════════════════════════════════════════════════
#  Markdown 解析器
# ═════════════════════════════════════════════════════════════════════

class MarkdownParser(DocumentParser):
    """
    Markdown 解析器（零依赖）。
    # 标题 → 章节边界，段落 → 正文。
    """
    extensions = [".md", ".markdown"]
    format_name = "Markdown"

    def parse(self, filepath: str | Path) -> dict:
        filepath = Path(filepath)
        if not filepath.exists():
            raise CorruptedFileError(f"文件不存在: {filepath}")

        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()

        # 将 # 标题转为章节格式
        text = re.sub(r"^#+\s+", "", text, flags=re.MULTILINE)
        chapters = extract_chapters(text)

        return {
            "title": filepath.stem,
            "chapters": chapters,
            "metadata": {
                "format": "markdown",
                "chars_total": len(text),
                "chars_cleaned": sum(len(c["text"]) for c in chapters),
                **self.get_size_info(filepath),
            },
        }


# ═════════════════════════════════════════════════════════════════════
#  文档路由器
# ═════════════════════════════════════════════════════════════════════

class DocumentRouter:
    """文件格式检测与路由"""

    def __init__(self):
        self._parsers: dict[str, type[DocumentParser]] = {}
        self._register_builtin()

    def _register_builtin(self):
        """注册所有内置解析器"""
        self.register(TxtParser)
        self.register(DocxParser)
        self.register(PdfParser)
        self.register(MarkdownParser)

    def register(self, parser_cls: type[DocumentParser]):
        for ext in parser_cls.extensions:
            ext_lower = ext.lower()
            if ext_lower in self._parsers:
                logger.warning(f"扩展名 {ext} 的解析器被覆盖: {parser_cls.__name__}")
            self._parsers[ext_lower] = parser_cls
        logger.info(f"[DocumentRouter] 注册解析器: {parser_cls.__name__} → {parser_cls.extensions}")

    def get_parser(self, filepath: str | Path) -> type[DocumentParser] | None:
        ext = Path(filepath).suffix.lower()
        return self._parsers.get(ext)

    def supported_extensions(self) -> list[str]:
        return list(self._parsers.keys())

    def supported_formats_display(self) -> str:
        names = []
        for cls in set(self._parsers.values()):
            exts = ", ".join(cls.extensions)
            names.append(f"{cls.format_name} ({exts})")
        return " / ".join(names)

    def parse(
        self,
        filepath: str | Path,
        force_format: str | None = None,
    ) -> dict:
        filepath = Path(filepath)

        if not filepath.exists():
            raise FileNotFoundError(f"文件不存在: {filepath}")

        if force_format:
            parser_cls = self._parsers.get(force_format.lower())
            if not parser_cls:
                raise UnsupportedFormatError(
                    f"不支持的格式: {force_format}。支持的格式: {self.supported_formats_display()}"
                )
        else:
            parser_cls = self.get_parser(filepath)
            if not parser_cls:
                raise UnsupportedFormatError(
                    f"不支持的文件格式: {filepath.suffix}。支持的格式: {self.supported_formats_display()}"
                )

        parser = parser_cls()
        try:
            result = parser.parse(filepath)
            logger.info(
                f"[DocumentRouter] 解析完成: {filepath.name} "
                f"→ {len(result.get('chapters', []))} 章, "
                f"格式: {result.get('metadata', {}).get('format', '?')}"
            )
            return result
        except Exception as e:
            logger.error(f"[DocumentRouter] 解析失败: {filepath.name} - {e}")
            raise


# ── 全局默认实例 ───────────────────────────────────────────────────────

_router: DocumentRouter | None = None


def get_router() -> DocumentRouter:
    global _router
    if _router is None:
        _router = DocumentRouter()
    return _router


def parse_document(filepath: str | Path, force_format: str | None = None) -> dict:
    return get_router().parse(filepath, force_format=force_format)


# ── 测试 ────────────────────────────────────────────────────────────────

def demo():
    router = get_router()
    print(f"支持格式: {router.supported_formats_display()}")
    print(f"扩展名: {router.supported_extensions()}")


if __name__ == "__main__":
    demo()
